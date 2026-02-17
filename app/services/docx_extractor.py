import io

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


def extract_docx(data: bytes) -> str:
    """Extract text from DOCX with heading hierarchy, tables, headers/footers."""
    doc = Document(io.BytesIO(data))
    parts: list[str] = []

    # --- Section headers ---
    for section in doc.sections[:1]:
        header_text = _paragraphs_text(section.header.paragraphs)
        if header_text:
            parts.append(header_text)

    # --- Body: paragraphs and tables in document order ---
    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = _format_paragraph(block)
            if text:
                parts.append(text)
        elif isinstance(block, Table):
            md = _table_to_markdown(block)
            if md:
                parts.append(md)

    # --- Section footers ---
    for section in doc.sections[:1]:
        footer_text = _paragraphs_text(section.footer.paragraphs)
        if footer_text:
            parts.append(footer_text)

    return "\n\n".join(parts)


# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------

def _iter_block_items(parent):
    """Yield Paragraph and Table objects in document order."""
    body = parent.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def _format_paragraph(para: Paragraph) -> str:
    """Format a paragraph, preserving heading levels and list markers."""
    text = para.text.strip()
    if not text:
        return ""

    style_name = (para.style.name or "").lower()

    # Headings → markdown
    if style_name.startswith("heading"):
        try:
            level = int(style_name.replace("heading", "").strip())
        except ValueError:
            level = 1
        return "#" * level + " " + text

    # List items → bullets
    numpr = para._element.find(qn("w:pPr"))
    if numpr is not None and numpr.find(qn("w:numPr")) is not None:
        return "- " + text

    return text


def _table_to_markdown(table: Table) -> str:
    """Convert a docx Table to a markdown table string."""
    rows: list[list[str]] = []
    for row in table.rows:
        rows.append([cell.text.strip().replace("\n", " ") for cell in row.cells])

    if not rows:
        return ""

    col_count = max(len(r) for r in rows)
    for row in rows:
        while len(row) < col_count:
            row.append("")

    lines: list[str] = []
    lines.append("| " + " | ".join(rows[0]) + " |")
    lines.append("| " + " | ".join("---" for _ in range(col_count)) + " |")
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


def _paragraphs_text(paragraphs) -> str:
    """Join non-empty paragraph texts."""
    return "\n".join(p.text.strip() for p in paragraphs if p.text.strip())
